#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
工具箱
功能1：视频批量重命名
功能2：筛选未绑账户
功能3：模板自动填写
"""

import os
import json
import datetime
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from tkinter.scrolledtext import ScrolledText
from openpyxl import Workbook, load_workbook

# 尝试导入Word处理库
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# 尝试导入拖拽支持库
try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
    HAS_DND = True
except ImportError:
    HAS_DND = False


# 支持的视频格式
VIDEO_EXTENSIONS = {
    '.mp4', '.avi', '.mkv', '.mov', '.wmv', '.flv', '.webm',
    '.m4v', '.mpg', '.mpeg', '.3gp', '.ts', '.mts', '.vob'
}

# 撤销记录文件名
UNDO_FILENAME = "_undo.xlsx"

# 模板配置文件
TEMPLATE_CONFIG_FILE = "templates.json"


class ToolboxApp:
    def __init__(self, root):
        self.root = root
        self.root.title("🧰 工具箱 Toolbox")
        self.root.geometry("900x700")
        self.root.minsize(800, 600)
        self.root.resizable(True, True)
        
        # 配置现代化主题和样式
        self.setup_styles()
        
        # 设置窗口背景色
        self.root.configure(bg='#f5f5f5')
        
        # 创建标签页
        self.notebook = ttk.Notebook(root, style='Custom.TNotebook')
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=12, pady=12)
        
        # 功能1：视频批量重命名
        self.tab1 = ttk.Frame(self.notebook, style='Card.TFrame')
        self.notebook.add(self.tab1, text="  📹 视频批量重命名  ")
        self.setup_video_rename_tab()
        
        # 功能2：筛选未绑账户
        self.tab2 = ttk.Frame(self.notebook, style='Card.TFrame')
        self.notebook.add(self.tab2, text="  🔍 筛选未绑账户  ")
        self.setup_filter_account_tab()
        
        # 功能3：模板自动填写
        self.tab3 = ttk.Frame(self.notebook, style='Card.TFrame')
        self.notebook.add(self.tab3, text="  📝 模板自动填写  ")
        self.setup_template_fill_tab()
        
        # 功能4：数值转换
        self.tab4 = ttk.Frame(self.notebook, style='Card.TFrame')
        self.notebook.add(self.tab4, text="  🔢 数值转换  ")
        self.setup_value_convert_tab()
        
        # 功能5：表格匹配
        self.tab5 = ttk.Frame(self.notebook, style='Card.TFrame')
        self.notebook.add(self.tab5, text="  📊 表格匹配  ")
        self.setup_table_match_tab()
    
    def setup_styles(self):
        """配置现代化UI样式"""
        style = ttk.Style()
        
        # 使用clam主题作为基础（更现代）
        style.theme_use('clam')
        
        # 主背景色
        bg_color = '#f5f5f5'
        card_bg = '#ffffff'
        accent_color = '#1976D2'
        accent_hover = '#1565C0'
        success_color = '#43A047'
        danger_color = '#E53935'
        
        # Notebook样式
        style.configure('Custom.TNotebook', background=bg_color, borderwidth=0)
        style.configure('Custom.TNotebook.Tab',
                       background='#e0e0e0',
                       foreground='#424242',
                       padding=[20, 10],
                       font=('微软雅黑', 10))
        style.map('Custom.TNotebook.Tab',
                 background=[('selected', card_bg)],
                 foreground=[('selected', accent_color)])
        
        # Frame样式
        style.configure('Card.TFrame', background=card_bg)
        style.configure('Main.TFrame', background=bg_color)
        
        # Label样式
        style.configure('Title.TLabel',
                       background=card_bg,
                       foreground='#212121',
                       font=('微软雅黑', 11, 'bold'))
        style.configure('Card.TLabel',
                       background=card_bg,
                       foreground='#424242',
                       font=('微软雅黑', 10))
        style.configure('Hint.TLabel',
                       background=card_bg,
                       foreground='#757575',
                       font=('微软雅黑', 9))
        style.configure('Success.TLabel',
                       background=card_bg,
                       foreground=success_color,
                       font=('微软雅黑', 10, 'bold'))
        
        # Button样式 - 主要按钮
        style.configure('Primary.TButton',
                       background=accent_color,
                       foreground='white',
                       borderwidth=0,
                       focuscolor='none',
                       padding=[15, 8],
                       font=('微软雅黑', 10))
        style.map('Primary.TButton',
                 background=[('active', accent_hover), ('pressed', accent_hover)])
        
        # Button样式 - 次要按钮
        style.configure('Secondary.TButton',
                       background='#e0e0e0',
                       foreground='#424242',
                       borderwidth=0,
                       focuscolor='none',
                       padding=[15, 8],
                       font=('微软雅黑', 10))
        style.map('Secondary.TButton',
                 background=[('active', '#bdbdbd'), ('pressed', '#9e9e9e')])
        
        # Button样式 - 成功按钮
        style.configure('Success.TButton',
                       background=success_color,
                       foreground='white',
                       borderwidth=0,
                       focuscolor='none',
                       padding=[15, 8],
                       font=('微软雅黑', 10))
        style.map('Success.TButton',
                 background=[('active', '#2E7D32'), ('pressed', '#2E7D32')])
        
        # Button样式 - 危险按钮
        style.configure('Danger.TButton',
                       background=danger_color,
                       foreground='white',
                       borderwidth=0,
                       focuscolor='none',
                       padding=[15, 8],
                       font=('微软雅黑', 10))
        style.map('Danger.TButton',
                 background=[('active', '#C62828'), ('pressed', '#C62828')])
        
        # LabelFrame样式
        style.configure('Card.TLabelframe',
                       background=card_bg,
                       borderwidth=1,
                       relief='solid')
        style.configure('Card.TLabelframe.Label',
                       background=card_bg,
                       foreground=accent_color,
                       font=('微软雅黑', 10, 'bold'))
        
        # Entry样式
        style.configure('Custom.TEntry',
                       fieldbackground='white',
                       borderwidth=1,
                       relief='solid',
                       padding=5)
        
        # Radiobutton样式
        style.configure('Custom.TRadiobutton',
                       background=card_bg,
                       foreground='#424242',
                       font=('微软雅黑', 10))
        
        # Checkbutton样式
        style.configure('Custom.TCheckbutton',
                       background=card_bg,
                       foreground='#424242',
                       font=('微软雅黑', 10))
    
    def paste_path(self, textvariable):
        """粘贴路径到输入框"""
        try:
            content = self.root.clipboard_get()
            content = content.strip('"').strip('{}')
            textvariable.set(content)
        except:
            pass

    def make_draggable(self, entry, textvariable, file_type='file', title='文件'):
        """为输入框添加拖拽支持
        file_type: 'file' (单个文件) 或 'folder' (文件夹) 或 'any' (任意)
        """
        if not HAS_DND:
            return

        def parse_drop_data(data):
            """解析拖拽数据，处理花括号包裹的路径（包含空格的路径）"""
            # 拖拽数据可能用花括号包裹包含空格的路径
            result = []
            i = 0
            data = data.strip()
            while i < len(data):
                if data[i] == '{':
                    # 找到匹配的右花括号
                    j = data.find('}', i)
                    if j != -1:
                        result.append(data[i+1:j])
                        i = j + 1
                    else:
                        i += 1
                elif data[i] == ' ':
                    i += 1
                else:
                    # 普通路径，读取到下一个空格或花括号
                    j = i
                    while j < len(data) and data[j] not in (' ', '{'):
                        j += 1
                    if j > i:
                        result.append(data[i:j])
                    i = j
            return result

        def on_drop(event):
            data = event.data
            paths = parse_drop_data(data)
            if not paths:
                return

            path = paths[0]
            # 根据类型验证
            if file_type == 'file':
                if not os.path.isfile(path):
                    messagebox.showwarning("提示", f"请拖入{title}文件\n当前：{path}")
                    return
            elif file_type == 'folder':
                if not os.path.isdir(path):
                    messagebox.showwarning("提示", f"请拖入{title}文件夹\n当前：{path}")
                    return
            # file_type == 'any' 不做验证

            textvariable.set(path)
            # 改变边框颜色表示已接收
            entry.config(bg='#e8f5e9')

        def on_drag_enter(event):
            entry.config(bg='#fff9c4')  # 拖入时变黄色

        def on_drag_leave(event):
            entry.config(bg='white')  # 离开时恢复

        # 注册拖拽事件
        entry.drop_target_register(DND_FILES)
        entry.dnd_bind('<<Drop>>', on_drop)
        entry.dnd_bind('<<DropEnter>>', on_drag_enter)
        entry.dnd_bind('<<DropLeave>>', on_drag_leave)
    
    # ==================== 功能1：视频批量重命名 ====================
    
    def setup_video_rename_tab(self):
        """设置视频重命名标签页"""
        main_frame = ttk.Frame(self.tab1, padding="15", style='Card.TFrame')
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        self.original_xlsx_path = tk.StringVar()  # 原始表格路径
        self.xlsx_path = tk.StringVar()
        self.folder_path = tk.StringVar()  # 视频文件夹（执行重命名时需要）
        self.processed_names = []  # 处理后的视频命名列表
        self.processed_data = []  # 处理后的完整数据
        self.video_files = []  # 实际扫描到的视频文件
        self.process_type = tk.StringVar(value="抖音")  # 整理方式：抖音/视频号
        
        # 标题
        ttk.Label(main_frame, text="视频批量重命名工具", style='Title.TLabel').pack(anchor='w', pady=(0, 10))
        
        # 整理方式选择
        row0 = ttk.Frame(main_frame, style='Card.TFrame')
        row0.pack(fill=tk.X, pady=5)
        ttk.Label(row0, text="整理方式：", style='Card.TLabel').pack(side=tk.LEFT)
        ttk.Radiobutton(row0, text="抖音", variable=self.process_type, value="抖音", style='Custom.TRadiobutton').pack(side=tk.LEFT, padx=10)
        ttk.Radiobutton(row0, text="视频号", variable=self.process_type, value="视频号", style='Custom.TRadiobutton').pack(side=tk.LEFT, padx=10)
        
        # 原始表格区域
        row1 = ttk.Frame(main_frame, style='Card.TFrame')
        row1.pack(fill=tk.X, pady=5)
        ttk.Label(row1, text="原始表格：", style='Card.TLabel', width=12).pack(side=tk.LEFT)
        entry1 = tk.Entry(row1, textvariable=self.original_xlsx_path, width=55, font=('微软雅黑', 9), relief='solid', bd=1)
        entry1.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        entry1.bind('<Double-Button-1>', lambda e: self.paste_path(self.original_xlsx_path))
        self.make_draggable(entry1, self.original_xlsx_path, 'file', '表格')
        ttk.Button(row1, text="选择", command=self.select_original_xlsx, style='Secondary.TButton').pack(side=tk.LEFT, padx=2)
        ttk.Button(row1, text="整理", command=self.process_original_xlsx, style='Primary.TButton').pack(side=tk.LEFT, padx=2)
        ttk.Button(row1, text="导出", command=self.export_processed_xlsx, style='Success.TButton').pack(side=tk.LEFT, padx=2)
        self.processed_label = ttk.Label(row1, text="已整理 0 条", style='Success.TLabel')
        self.processed_label.pack(side=tk.LEFT, padx=10)
        
        # 视频文件夹区域
        row2 = ttk.Frame(main_frame, style='Card.TFrame')
        row2.pack(fill=tk.X, pady=5)
        ttk.Label(row2, text="视频文件夹：", style='Card.TLabel', width=12).pack(side=tk.LEFT)
        entry2 = tk.Entry(row2, textvariable=self.folder_path, width=55, font=('微软雅黑', 9), relief='solid', bd=1)
        entry2.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        entry2.bind('<Double-Button-1>', lambda e: self.paste_path(self.folder_path))
        self.make_draggable(entry2, self.folder_path, 'folder', '视频文件夹')
        ttk.Button(row2, text="选择", command=self.select_folder, style='Secondary.TButton').pack(side=tk.LEFT, padx=2)
        ttk.Button(row2, text="扫描核对", command=self.scan_and_check, style='Primary.TButton').pack(side=tk.LEFT, padx=2)
        self.file_count_label = ttk.Label(row2, text="共 0 个视频", style='Hint.TLabel')
        self.file_count_label.pack(side=tk.LEFT, padx=10)
        
        # 重命名表格区域（自动生成，无需手动操作）
        row3 = ttk.Frame(main_frame, style='Card.TFrame')
        row3.pack(fill=tk.X, pady=5)
        ttk.Label(row3, text="重命名表格：", style='Card.TLabel', width=12).pack(side=tk.LEFT)
        self.rename_table_label = ttk.Label(row3, text="整理后自动生成", style='Hint.TLabel')
        self.rename_table_label.pack(side=tk.LEFT, padx=5)
        
        # 操作区域
        row4 = ttk.Frame(main_frame, style='Card.TFrame')
        row4.pack(fill=tk.X, pady=10)
        ttk.Label(row4, text="操作：", style='Card.TLabel', width=12).pack(side=tk.LEFT)
        ttk.Button(row4, text="📋 预览", command=self.preview_rename, style='Secondary.TButton').pack(side=tk.LEFT, padx=5)
        ttk.Button(row4, text="✅ 执行", command=self.execute_rename, style='Success.TButton').pack(side=tk.LEFT, padx=5)
        ttk.Button(row4, text="↩️ 撤销", command=self.undo_rename, style='Danger.TButton').pack(side=tk.LEFT, padx=5)
        
        # 日志区域
        log_frame = ttk.LabelFrame(main_frame, text="操作日志", padding="10", style='Card.TLabelframe')
        log_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        self.log_text1 = ScrolledText(log_frame, height=15, font=('Consolas', 9), bg='#fafafa', relief='flat', bd=0)
        self.log_text1.pack(fill=tk.BOTH, expand=True)
    
    def log1(self, msg):
        self.log_text1.insert(tk.END, msg + "\n")
        self.log_text1.see(tk.END)
    
    def clear_log1(self):
        self.log_text1.delete(1.0, tk.END)
    
    def select_original_xlsx(self):
        """选择原始表格"""
        file = filedialog.askopenfilename(title="选择原始表格", filetypes=[("Excel", "*.xlsx"), ("Excel", "*.xls"), ("CSV", "*.csv")])
        if file:
            self.original_xlsx_path.set(file)
            self.log1(f"原始表格：{file}")
    
    def process_original_xlsx(self):
        """整理原始表格，根据整理方式（抖音/视频号）应用不同规则"""
        file = self.original_xlsx_path.get()
        if not file:
            messagebox.showwarning("提示", "请先选择原始表格")
            return
        
        if not os.path.exists(file):
            messagebox.showwarning("提示", "文件不存在")
            return
        
        process_type = self.process_type.get()
        self.processed_names = []
        self.processed_data = []
        self.clear_log1()
        
        try:
            self.log1(f"整理方式：{process_type}")
            self.log1(f"读取文件：{file}")
            
            # 根据文件类型选择处理方式
            if file.lower().endswith('.csv'):
                self._process_csv_file(file, process_type)
            else:
                wb = load_workbook(file, data_only=True)
                ws = wb.active
                max_row = min(ws.max_row, 301)
                
                if process_type == "抖音":
                    self._process_douyin(ws, max_row)
                else:
                    self._process_weixin(ws, max_row)
                
                wb.close()
            
            # 显示整理结果
            self.log1("")
            name_col = '视频标题' if process_type == "抖音" else '视频名称'
            for i, data in enumerate(self.processed_data[:10]):
                self.log1(f"  {i+1}. {data[name_col]} → {data['视频命名']}")
            if len(self.processed_data) > 10:
                self.log1(f"  ... 共 {len(self.processed_data)} 条")
            
            self.processed_label.config(text=f"已整理 {len(self.processed_names)} 条")
            cost_name = '消耗' if process_type == "抖音" else '花费'
            self.log1(f"\n共整理 {len(self.processed_names)} 条数据（筛选{cost_name}>1500）")
            
            # 自动生成重命名表格
            self._auto_generate_rename_table(file, process_type)
            
            messagebox.showinfo("完成", f"整理完成，共 {len(self.processed_names)} 条\n重命名表格已自动生成")
            
        except Exception as e:
            messagebox.showerror("错误", str(e))
            self.log1(f"错误：{e}")
    
    def _process_csv_file(self, file, process_type):
        """处理CSV文件"""
        import csv
        
        self.log1("文件格式：CSV")
        
        try:
            # 尝试不同编码
            encodings = ['utf-8', 'gbk', 'gb18030', 'utf-8-sig']
            rows = None
            
            for encoding in encodings:
                try:
                    with open(file, 'r', encoding=encoding) as f:
                        reader = csv.reader(f)
                        rows = list(reader)
                    break
                except:
                    continue
            
            if rows is None:
                messagebox.showerror("错误", "无法读取CSV文件，请检查文件编码")
                return
            
            max_row = min(len(rows), 301)  # 最多300行
            
            if process_type == "抖音":
                self._process_douyin_csv(rows, max_row)
            else:
                self._process_weixin_csv(rows, max_row)
            
        except Exception as e:
            messagebox.showerror("错误", f"CSV处理错误：{e}")
            self.log1(f"错误：{e}")
    
    def _process_douyin_csv(self, rows, max_row):
        """抖音规则处理CSV"""
        self.log1("抖音规则：")
        self.log1("  - 保留列：视频标题、消耗")
        self.log1("  - 筛选：消耗>1500")
        
        if len(rows) < 1:
            messagebox.showwarning("提示", "CSV文件没有数据")
            return
        
        # 查找列
        header_row = rows[0]
        title_col = None
        cost_col = None
        
        for i, h in enumerate(header_row):
            if h and '视频标题' in str(h):
                title_col = i
            if h and '消耗' in str(h):
                cost_col = i
        
        if title_col is None or cost_col is None:
            messagebox.showwarning("提示", "CSV中未找到'视频标题'或'消耗'列")
            return
        
        self.log1(f"视频标题列：第{title_col+1}列")
        self.log1(f"消耗列：第{cost_col+1}列")
        
        for row_num in range(1, max_row):
            row_data = rows[row_num]
            if len(row_data) <= max(title_col, cost_col):
                continue
            
            title = row_data[title_col] if title_col < len(row_data) else None
            cost_raw = row_data[cost_col] if cost_col < len(row_data) else None
            
            if not title:
                continue
            
            title = str(title).strip()
            
            # 处理消耗值
            cost_value = None
            cost_str = ''
            
            if cost_raw is not None:
                cost_str = str(cost_raw).strip()
                cost_str_clean = cost_str.replace(',', '').replace(' ', '').replace('￥', '').replace('¥', '')
                try:
                    cost_value = float(cost_str_clean)
                except:
                    continue
            
            # 筛选：消耗必须大于1500
            if cost_value is None or cost_value <= 1500:
                continue
            
            # 视频命名
            name2 = title.replace('推送视频_', '').replace('-衍生', '')
            video_name = f"消耗{cost_str}-{name2}"
            
            self.processed_names.append(video_name)
            self.processed_data.append({
                '视频标题': title,
                '消耗': cost_str,
                '视频命名2': name2,
                '视频命名': video_name
            })
    
    def _process_weixin_csv(self, rows, max_row):
        """视频号规则处理CSV"""
        self.log1("视频号规则：")
        self.log1("  - 保留列：视频名称、花费、下单次数、下单金额、下单ROI")
        self.log1("  - 筛选：花费>1500")
        
        if len(rows) < 1:
            messagebox.showwarning("提示", "CSV文件没有数据")
            return
        
        # 查找列
        header_row = rows[0]
        name_col = None
        cost_col = None
        order_count_col = None
        order_amount_col = None
        roi_col = None
        
        for i, h in enumerate(header_row):
            if h:
                h_str = str(h)
                if '视频名称' in h_str:
                    name_col = i
                elif '花费' in h_str:
                    cost_col = i
                elif '下单次数' in h_str:
                    order_count_col = i
                elif '下单金额' in h_str:
                    order_amount_col = i
                elif '下单ROI' in h_str or 'ROI' in h_str:
                    roi_col = i
        
        if name_col is None or cost_col is None or order_count_col is None:
            messagebox.showwarning("提示", "CSV中未找到'视频名称'、'花费'或'下单次数'列")
            return
        
        self.log1(f"视频名称列：第{name_col+1}列")
        self.log1(f"花费列：第{cost_col+1}列")
        self.log1(f"下单次数列：第{order_count_col+1}列")
        
        for row_num in range(1, max_row):
            row_data = rows[row_num]
            if len(row_data) <= max(name_col, cost_col, order_count_col):
                continue
            
            name = row_data[name_col] if name_col < len(row_data) else None
            cost_raw = row_data[cost_col] if cost_col < len(row_data) else None
            order_count_raw = row_data[order_count_col] if order_count_col < len(row_data) else None
            order_amount_raw = row_data[order_amount_col] if order_amount_col and order_amount_col < len(row_data) else None
            roi_raw = row_data[roi_col] if roi_col and roi_col < len(row_data) else None
            
            if not name:
                continue
            
            name = str(name).strip()
            
            # 处理花费值
            cost_value = None
            cost_str = ''
            
            if cost_raw is not None:
                cost_str = str(cost_raw).strip()
                cost_str_clean = cost_str.replace(',', '').replace(' ', '').replace('￥', '').replace('¥', '')
                try:
                    cost_value = float(cost_str_clean)
                except:
                    continue
            
            # 筛选：花费必须大于1500
            if cost_value is None or cost_value <= 1500:
                continue
            
            # 处理下单次数
            order_count_str = str(order_count_raw).strip() if order_count_raw else ''
            
            # 处理下单金额
            order_amount_str = str(order_amount_raw).strip() if order_amount_raw else ''
            
            # 处理ROI
            roi_str = str(roi_raw).strip() if roi_raw else ''
            
            # 视频命名
            video_name = f"转{order_count_str}-{name}"
            
            self.processed_names.append(video_name)
            self.processed_data.append({
                '视频名称': name,
                '花费': cost_str,
                '下单次数': order_count_str,
                '下单金额': order_amount_str,
                '下单ROI': roi_str,
                '视频命名': video_name
            })
    
    def _process_douyin(self, ws, max_row):
        """抖音规则：只保留视频标题、消耗；筛选消耗>1500"""
        self.log1("抖音规则：")
        self.log1("  - 保留列：视频标题、消耗")
        self.log1("  - 筛选：消耗>1500")
        self.log1("  - 视频命名：消耗{值}-{视频标题去掉推送视频_和-衍生}")
        
        # 查找列
        header_row = [cell.value for cell in ws[1]]
        title_col = None
        cost_col = None
        
        for i, h in enumerate(header_row):
            if h and '视频标题' in str(h):
                title_col = i + 1
            if h and '消耗' in str(h):
                cost_col = i + 1
        
        if not title_col or not cost_col:
            messagebox.showwarning("提示", "表格中未找到'视频标题'或'消耗'列")
            return
        
        self.log1(f"视频标题列：第{title_col}列")
        self.log1(f"消耗列：第{cost_col}列")
        
        for row_num in range(2, max_row + 1):
            title = ws.cell(row=row_num, column=title_col).value
            cost_raw = ws.cell(row=row_num, column=cost_col).value
            
            if not title:
                continue
            
            title = str(title).strip()
            
            # 处理消耗值
            cost_value = None
            cost_str = ''
            
            if cost_raw is not None:
                if isinstance(cost_raw, (int, float)):
                    cost_value = float(cost_raw)
                    cost_str = str(int(cost_raw) if cost_raw == int(cost_raw) else cost_raw)
                else:
                    cost_str = str(cost_raw).strip()
                    cost_str_clean = cost_str.replace(',', '').replace(' ', '').replace('￥', '').replace('¥', '')
                    try:
                        cost_value = float(cost_str_clean)
                    except:
                        continue
            
            # 筛选：消耗必须大于1500
            if cost_value is None or cost_value <= 1500:
                continue
            
            # 视频命名 = "消耗" + 消耗值 + "-" + (视频标题去掉"推送视频_"和"-衍生")
            name2 = title.replace('推送视频_', '').replace('-衍生', '')
            video_name = f"消耗{cost_str}-{name2}"
            
            self.processed_names.append(video_name)
            self.processed_data.append({
                '视频标题': title,
                '消耗': cost_str,
                '视频命名2': name2,
                '视频命名': video_name
            })
    
    def _process_weixin(self, ws, max_row):
        """视频号规则：只保留视频名称、花费、下单次数、下单金额、下单ROI；筛选花费>1500"""
        self.log1("视频号规则：")
        self.log1("  - 保留列：视频名称、花费、下单次数、下单金额、下单ROI")
        self.log1("  - 筛选：花费>1500")
        self.log1("  - 视频命名：转{下单次数}-{视频名称}")
        
        # 查找列
        header_row = [cell.value for cell in ws[1]]
        name_col = None
        cost_col = None
        order_count_col = None
        order_amount_col = None
        roi_col = None
        
        for i, h in enumerate(header_row):
            if h:
                h_str = str(h)
                if '视频名称' in h_str:
                    name_col = i + 1
                elif '花费' in h_str:
                    cost_col = i + 1
                elif '下单次数' in h_str:
                    order_count_col = i + 1
                elif '下单金额' in h_str:
                    order_amount_col = i + 1
                elif '下单ROI' in h_str or 'ROI' in h_str:
                    roi_col = i + 1
        
        if not name_col or not cost_col or not order_count_col:
            messagebox.showwarning("提示", "表格中未找到'视频名称'、'花费'或'下单次数'列")
            return
        
        self.log1(f"视频名称列：第{name_col}列")
        self.log1(f"花费列：第{cost_col}列")
        self.log1(f"下单次数列：第{order_count_col}列")
        
        for row_num in range(2, max_row + 1):
            name = ws.cell(row=row_num, column=name_col).value
            cost_raw = ws.cell(row=row_num, column=cost_col).value
            order_count_raw = ws.cell(row=row_num, column=order_count_col).value
            order_amount_raw = ws.cell(row=row_num, column=order_amount_col).value if order_amount_col else None
            roi_raw = ws.cell(row=row_num, column=roi_col).value if roi_col else None
            
            if not name:
                continue
            
            name = str(name).strip()
            
            # 处理花费值
            cost_value = None
            cost_str = ''
            
            if cost_raw is not None:
                if isinstance(cost_raw, (int, float)):
                    cost_value = float(cost_raw)
                    cost_str = str(int(cost_raw) if cost_raw == int(cost_raw) else cost_raw)
                else:
                    cost_str = str(cost_raw).strip()
                    cost_str_clean = cost_str.replace(',', '').replace(' ', '').replace('￥', '').replace('¥', '')
                    try:
                        cost_value = float(cost_str_clean)
                    except:
                        continue
            
            # 筛选：花费必须大于1500
            if cost_value is None or cost_value <= 1500:
                continue
            
            # 处理下单次数
            order_count_str = ''
            if order_count_raw is not None:
                if isinstance(order_count_raw, (int, float)):
                    order_count_str = str(int(order_count_raw) if order_count_raw == int(order_count_raw) else order_count_raw)
                else:
                    order_count_str = str(order_count_raw).strip()
            
            # 处理下单金额
            order_amount_str = ''
            if order_amount_raw is not None:
                if isinstance(order_amount_raw, (int, float)):
                    order_amount_str = str(int(order_amount_raw) if order_amount_raw == int(order_amount_raw) else order_amount_raw)
                else:
                    order_amount_str = str(order_amount_raw).strip()
            
            # 处理ROI
            roi_str = ''
            if roi_raw is not None:
                if isinstance(roi_raw, (int, float)):
                    roi_str = str(roi_raw)
                else:
                    roi_str = str(roi_raw).strip()
            
            # 视频命名 = "转" + 下单次数 + "-" + 视频名称
            video_name = f"转{order_count_str}-{name}"
            
            self.processed_names.append(video_name)
            self.processed_data.append({
                '视频名称': name,
                '花费': cost_str,
                '下单次数': order_count_str,
                '下单金额': order_amount_str,
                '下单ROI': roi_str,
                '视频命名': video_name
            })
    
    def export_processed_xlsx(self):
        """导出整理后的数据和重命名表格到一个工作簿（两个工作表）"""
        if not self.processed_data:
            messagebox.showwarning("提示", "请先整理原始表格")
            return

        process_type = self.process_type.get()

        file = filedialog.asksaveasfilename(
            title="保存整理结果", defaultextension=".xlsx",
            filetypes=[("Excel", "*.xlsx")], initialfile=f"{process_type}整理结果.xlsx")
        if not file:
            return

        try:
            wb = Workbook()
            # 工作表1：整理结果
            ws1 = wb.active
            ws1.title = "整理结果"

            if process_type == "抖音":
                ws1['A1'] = '视频标题'
                ws1['B1'] = '消耗'
                ws1['C1'] = '视频命名2'
                ws1['D1'] = '视频命名'

                for i, data in enumerate(self.processed_data):
                    row = i + 2
                    ws1[f'A{row}'] = data['视频标题']
                    ws1[f'B{row}'] = data['消耗']
                    ws1[f'C{row}'] = data['视频命名2']
                    ws1[f'D{row}'] = data['视频命名']

                ws1.column_dimensions['A'].width = 40
                ws1.column_dimensions['B'].width = 15
                ws1.column_dimensions['C'].width = 30
                ws1.column_dimensions['D'].width = 50
            else:
                # 视频号
                ws1['A1'] = '视频名称'
                ws1['B1'] = '花费'
                ws1['C1'] = '下单次数'
                ws1['D1'] = '下单金额'
                ws1['E1'] = '下单ROI'
                ws1['F1'] = '视频命名'

                for i, data in enumerate(self.processed_data):
                    row = i + 2
                    ws1[f'A{row}'] = data['视频名称']
                    ws1[f'B{row}'] = data['花费']
                    ws1[f'C{row}'] = data['下单次数']
                    ws1[f'D{row}'] = data['下单金额']
                    ws1[f'E{row}'] = data['下单ROI']
                    ws1[f'F{row}'] = data['视频命名']

                ws1.column_dimensions['A'].width = 40
                ws1.column_dimensions['B'].width = 15
                ws1.column_dimensions['C'].width = 15
                ws1.column_dimensions['D'].width = 15
                ws1.column_dimensions['E'].width = 15
                ws1.column_dimensions['F'].width = 50

            # 工作表2：重命名表格
            ws2 = wb.create_sheet(title="重命名表格")
            ws2['A1'] = '文件名（旧）'
            ws2['B1'] = '文件名（新）'

            for i, new_name in enumerate(self.processed_names):
                row = i + 2

                # 根据整理方式生成旧文件名
                if process_type == "抖音":
                    if i == 0:
                        old_name = "巨量引擎工作台-升级版.mp4"
                    else:
                        old_name = f"巨量引擎工作台-升级版 ({i}).mp4"
                else:
                    # 视频号
                    if i == 0:
                        old_name = "腾讯营销_-_客户工作台.mp4"
                    else:
                        old_name = f"腾讯营销_-_客户工作台 ({i}).mp4"

                ws2[f'A{row}'] = old_name
                ws2[f'B{row}'] = new_name + ".mp4"

            ws2.column_dimensions['A'].width = 35
            ws2.column_dimensions['B'].width = 50

            wb.save(file)

            # 同时更新重命名表格路径（供后续预览/执行使用）
            self.xlsx_path.set(file)

            self.log1(f"已导出（含2个工作表）：{file}")
            self.log1(f"  工作表1：整理结果（{len(self.processed_data)} 条）")
            self.log1(f"  工作表2：重命名表格（{len(self.processed_names)} 条）")
            messagebox.showinfo("成功", f"导出完成\n包含2个工作表：\n1. 整理结果\n2. 重命名表格\n\n文件：{file}")

        except Exception as e:
            messagebox.showerror("错误", str(e))
            self.log1(f"错误：{e}")
    
    def select_folder(self):
        folder = filedialog.askdirectory(title="选择视频文件夹")
        if folder:
            self.folder_path.set(folder)
            self.log1(f"文件夹：{folder}")
    
    def scan_and_check(self):
        """扫描视频文件夹并核对生成的旧文件名是否存在"""
        folder = self.folder_path.get()
        if not folder:
            messagebox.showwarning("提示", "请先选择视频文件夹")
            return
        
        if not self.processed_names:
            messagebox.showwarning("提示", "请先整理原始表格")
            return
        
        process_type = self.process_type.get()
        
        # 扫描视频文件
        self.video_files = []
        try:
            for f in os.listdir(folder):
                if os.path.isfile(os.path.join(folder, f)):
                    ext = os.path.splitext(f)[1].lower()
                    if ext in VIDEO_EXTENSIONS:
                        self.video_files.append(f)
            
            self.file_count_label.config(text=f"共 {len(self.video_files)} 个")
            
            # 生成预期的旧文件名列表
            expected_names = []
            for i in range(len(self.processed_names)):
                if process_type == "抖音":
                    if i == 0:
                        expected_names.append("巨量引擎工作台-升级版.mp4")
                    else:
                        expected_names.append(f"巨量引擎工作台-升级版 ({i}).mp4")
                else:
                    # 视频号
                    if i == 0:
                        expected_names.append("腾讯营销_-_客户工作台.mp4")
                    else:
                        expected_names.append(f"腾讯营销_-_客户工作台 ({i}).mp4")
            
            # 核对：检查每个预期文件名是否存在于实际文件列表中
            self.log1(f"扫描到 {len(self.video_files)} 个视频文件")
            self.log1(f"整理数据 {len(self.processed_names)} 条")
            self.log1("")
            
            matched = 0
            missing = []
            
            for i, expected in enumerate(expected_names):
                if expected in self.video_files:
                    matched += 1
                else:
                    missing.append(expected)
            
            self.log1(f"匹配结果：")
            self.log1(f"  ✓ 找到 {matched} 个文件")
            if missing:
                self.log1(f"  ✗ 缺失 {len(missing)} 个文件：")
                for m in missing:
                    self.log1(f"      {m}")
            
            self.log1("")
            
            if len(self.video_files) != len(self.processed_names):
                self.log1(f"⚠ 数量不一致！视频文件:{len(self.video_files)} vs 整理数据:{len(self.processed_names)}")
            
            if matched == len(expected_names) and len(self.video_files) == len(self.processed_names):
                self.log1("✓ 所有文件匹配，可以执行重命名")
                messagebox.showinfo("核对完成", f"全部匹配！\n视频文件: {len(self.video_files)} 个\n整理数据: {len(self.processed_names)} 条")
            else:
                self.log1(f"⚠ 存在问题，请检查")
                messagebox.showwarning("核对完成", f"存在问题！\n找到: {matched}\n缺失: {len(missing)}")
            
        except Exception as e:
            messagebox.showerror("错误", str(e))
            self.log1(f"错误：{e}")
    
    def _auto_generate_rename_table(self, original_file, process_type):
        """整理完成后自动生成重命名表格，保存到原始表格同目录"""
        if not self.processed_names:
            return

        # 生成保存路径：原始表格同目录
        dir_path = os.path.dirname(original_file)
        file_name = f"{process_type}视频重命名.xlsx"
        save_path = os.path.join(dir_path, file_name)

        try:
            wb = Workbook()
            ws = wb.active
            ws['A1'] = '文件名（旧）'
            ws['B1'] = '文件名（新）'

            # 生成旧文件名
            for i, new_name in enumerate(self.processed_names):
                row = i + 2

                # 根据整理方式生成旧文件名
                if process_type == "抖音":
                    if i == 0:
                        old_name = "巨量引擎工作台-升级版.mp4"
                    else:
                        old_name = f"巨量引擎工作台-升级版 ({i}).mp4"
                else:
                    # 视频号
                    if i == 0:
                        old_name = "腾讯营销_-_客户工作台.mp4"
                    else:
                        old_name = f"腾讯营销_-_客户工作台 ({i}).mp4"

                ws[f'A{row}'] = old_name
                ws[f'B{row}'] = new_name + ".mp4"

            ws.column_dimensions['A'].width = 35
            ws.column_dimensions['B'].width = 50
            wb.save(save_path)
            self.xlsx_path.set(save_path)

            # 更新界面显示
            self.rename_table_label.config(
                text=f"已生成：{file_name}",
                style='Success.TLabel'
            )

            self.log1(f"重命名表格已自动生成：{save_path}")

        except Exception as e:
            messagebox.showerror("错误", f"生成重命名表格失败：{e}")
            self.log1(f"错误：{e}")

    def export_xlsx(self):
        """导出重命名表格，自动生成旧文件名并填充新文件名"""
        if not self.processed_names:
            messagebox.showwarning("提示", "请先整理原始表格")
            return
        
        process_type = self.process_type.get()
        
        file = filedialog.asksaveasfilename(
            title="保存重命名表格", defaultextension=".xlsx",
            filetypes=[("Excel", "*.xlsx")], initialfile=f"{process_type}视频重命名.xlsx")
        if not file:
            return
        
        try:
            wb = Workbook()
            ws = wb.active
            ws['A1'] = '文件名（旧）'
            ws['B1'] = '文件名（新）'
            
            # 生成旧文件名
            for i, new_name in enumerate(self.processed_names):
                row = i + 2
                
                # 根据整理方式生成旧文件名
                if process_type == "抖音":
                    if i == 0:
                        old_name = "巨量引擎工作台-升级版.mp4"
                    else:
                        old_name = f"巨量引擎工作台-升级版 ({i}).mp4"
                else:
                    # 视频号
                    if i == 0:
                        old_name = "腾讯营销_-_客户工作台.mp4"
                    else:
                        old_name = f"腾讯营销_-_客户工作台 ({i}).mp4"
                
                ws[f'A{row}'] = old_name
                ws[f'B{row}'] = new_name + ".mp4"
            
            ws.column_dimensions['A'].width = 35
            ws.column_dimensions['B'].width = 50
            wb.save(file)
            self.xlsx_path.set(file)
            
            self.log1(f"已导出：{file}")
            self.log1(f"共 {len(self.processed_names)} 条重命名记录")
            messagebox.showinfo("成功", f"导出完成\n共 {len(self.processed_names)} 条")
            
        except Exception as e:
            messagebox.showerror("错误", str(e))
            self.log1(f"错误：{e}")
    
    def select_xlsx(self):
        file = filedialog.askopenfilename(title="选择表格", filetypes=[("Excel", "*.xlsx")])
        if file:
            self.xlsx_path.set(file)
            self.log1(f"已加载：{file}")
    
    def load_xlsx_data(self):
        file = self.xlsx_path.get()
        if not file or not os.path.isfile(file):
            return None

        try:
            wb = load_workbook(file, data_only=True)
            # 优先读取"重命名表格"工作表，没有则读取活动工作表
            if "重命名表格" in wb.sheetnames:
                ws = wb["重命名表格"]
            else:
                ws = wb.active
            result = []
            for row in ws.iter_rows(min_row=2, values_only=True):
                if len(row) >= 2 and row[0] and row[1]:
                    result.append((str(row[0]).strip(), str(row[1]).strip()))
            wb.close()
            return result if result else None
        except Exception as e:
            messagebox.showerror("错误", str(e))
            return None
    
    def preview_rename(self):
        data = self.load_xlsx_data()
        if not data:
            messagebox.showwarning("提示", "请先加载表格")
            return
        
        folder = self.folder_path.get()
        if not folder:
            messagebox.showwarning("提示", "请先选择文件夹")
            return
        
        self.clear_log1()
        ok, err = 0, 0
        for old, new in data:
            if os.path.exists(os.path.join(folder, old)):
                self.log1(f"[OK] {old} -> {new}")
                ok += 1
            else:
                self.log1(f"[不存在] {old}")
                err += 1
        self.log1(f"总计 {len(data)}，可用 {ok}，问题 {err}")
    
    def save_undo(self, folder, data):
        file = os.path.join(folder, UNDO_FILENAME)
        wb = Workbook()
        ws = wb.active
        ws['A1'] = '当前'
        ws['B1'] = '原名'
        for i, (old, new) in enumerate(data, 2):
            ws[f'A{i}'] = new
            ws[f'B{i}'] = old
        wb.save(file)
    
    def execute_rename(self):
        data = self.load_xlsx_data()
        if not data:
            return
        
        folder = self.folder_path.get()
        if not folder:
            return
        
        if not messagebox.askyesno("确认", f"确定重命名 {len(data)} 个文件？"):
            return
        
        self.clear_log1()
        ok, err, success = 0, 0, []
        for old, new in data:
            old_path = os.path.join(folder, old)
            new_path = os.path.join(folder, new)
            try:
                if not os.path.exists(old_path):
                    self.log1(f"[失败] 不存在：{old}")
                    err += 1
                elif os.path.exists(new_path):
                    self.log1(f"[跳过] 已存在：{new}")
                    err += 1
                else:
                    os.rename(old_path, new_path)
                    self.log1(f"[成功] {old} -> {new}")
                    ok += 1
                    success.append((old, new))
            except Exception as e:
                self.log1(f"[失败] {old}：{e}")
                err += 1
        
        if success:
            self.save_undo(folder, success)
        self.log1(f"完成！成功 {ok}，失败 {err}")
        messagebox.showinfo("完成", f"成功 {ok}，失败 {err}")
    
    def undo_rename(self):
        folder = self.folder_path.get()
        if not folder:
            return
        
        undo_file = os.path.join(folder, UNDO_FILENAME)
        if not os.path.exists(undo_file):
            messagebox.showwarning("提示", "无撤销记录")
            return
        
        wb = load_workbook(undo_file, data_only=True)
        ws = wb.active
        data = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            if len(row) >= 2 and row[0] and row[1]:
                data.append((str(row[0]).strip(), str(row[1]).strip()))
        wb.close()
        
        if not messagebox.askyesno("确认", f"撤销 {len(data)} 个文件？"):
            return
        
        self.clear_log1()
        ok, err = 0, 0
        for cur, orig in data:
            try:
                if os.path.exists(os.path.join(folder, cur)):
                    os.rename(os.path.join(folder, cur), os.path.join(folder, orig))
                    self.log1(f"[成功] {cur} -> {orig}")
                    ok += 1
                else:
                    self.log1(f"[失败] 不存在：{cur}")
                    err += 1
            except Exception as e:
                self.log1(f"[失败] {e}")
                err += 1
        
        if ok > 0:
            os.remove(undo_file)
        self.log1(f"撤销完成！成功 {ok}")
    
    # ==================== 功能2：筛选未绑账户 ====================
    
    def setup_filter_account_tab(self):
        main_frame = ttk.Frame(self.tab2, padding="15", style='Card.TFrame')
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        self.account_xlsx = tk.StringVar()
        self.account_txt = tk.StringVar()
        self.unbound_accounts = []
        
        # 标题
        ttk.Label(main_frame, text="筛选未绑账户", style='Title.TLabel').pack(anchor='w', pady=(0, 10))
        
        # 表格文件
        row1 = ttk.Frame(main_frame, style='Card.TFrame')
        row1.pack(fill=tk.X, pady=5)
        ttk.Label(row1, text="表格文件：", style='Card.TLabel', width=12).pack(side=tk.LEFT)
        e1 = tk.Entry(row1, textvariable=self.account_xlsx, width=50, font=('微软雅黑', 9), relief='solid', bd=1)
        e1.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        e1.bind('<Double-Button-1>', lambda e: self.paste_path(self.account_xlsx))
        self.make_draggable(e1, self.account_xlsx, 'file', '表格')
        ttk.Button(row1, text="选择", command=self.select_account_xlsx, style='Secondary.TButton').pack(side=tk.LEFT, padx=2)
        ttk.Label(row1, text="(已绑定)", style='Hint.TLabel').pack(side=tk.LEFT, padx=5)
        
        # TXT文件
        row2 = ttk.Frame(main_frame, style='Card.TFrame')
        row2.pack(fill=tk.X, pady=5)
        ttk.Label(row2, text="TXT文件：", style='Card.TLabel', width=12).pack(side=tk.LEFT)
        e2 = tk.Entry(row2, textvariable=self.account_txt, width=50, font=('微软雅黑', 9), relief='solid', bd=1)
        e2.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        e2.bind('<Double-Button-1>', lambda e: self.paste_path(self.account_txt))
        self.make_draggable(e2, self.account_txt, 'file', 'TXT')
        ttk.Button(row2, text="选择", command=self.select_account_txt, style='Secondary.TButton').pack(side=tk.LEFT, padx=2)
        ttk.Label(row2, text="(所有账户)", style='Hint.TLabel').pack(side=tk.LEFT, padx=5)
        
        # 操作按钮
        row3 = ttk.Frame(main_frame, style='Card.TFrame')
        row3.pack(fill=tk.X, pady=10)
        ttk.Button(row3, text="🔍 开始筛选", command=self.filter_accounts, style='Primary.TButton').pack(side=tk.LEFT, padx=5)
        ttk.Button(row3, text="📋 复制结果", command=self.copy_accounts, style='Success.TButton').pack(side=tk.LEFT, padx=5)
        self.stats_label = ttk.Label(row3, text="", style='Success.TLabel')
        self.stats_label.pack(side=tk.LEFT, padx=20)
        
        # 结果区域
        result_frame = ttk.LabelFrame(main_frame, text="筛选结果", padding="10", style='Card.TLabelframe')
        result_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        self.log_text2 = ScrolledText(result_frame, height=20, font=('Consolas', 9), bg='#fafafa', relief='flat', bd=0)
        self.log_text2.pack(fill=tk.BOTH, expand=True)
    
    def log2(self, msg):
        self.log_text2.insert(tk.END, msg + "\n")
        self.log_text2.see(tk.END)
    
    def clear_log2(self):
        self.log_text2.delete(1.0, tk.END)
    
    def select_account_xlsx(self):
        file = filedialog.askopenfilename(title="选择表格", filetypes=[("Excel", "*.xlsx"), ("全部", "*.*")])
        if file:
            self.account_xlsx.set(file)
            self.log2(f"表格：{file}")
    
    def select_account_txt(self):
        file = filedialog.askopenfilename(title="选择TXT", filetypes=[("文本", "*.txt"), ("全部", "*.*")])
        if file:
            self.account_txt.set(file)
            self.log2(f"TXT：{file}")
    
    def load_bound(self):
        file = self.account_xlsx.get()
        if not file:
            return None
        wb = load_workbook(file, data_only=True)
        ws = wb.active
        accounts = set()
        for row in ws.iter_rows(min_row=1, values_only=True):
            if row and row[0]:
                accounts.add(str(row[0]).strip())
        wb.close()
        return accounts
    
    def load_all(self):
        file = self.account_txt.get()
        if not file:
            return None
        accounts = set()
        for enc in ['utf-8', 'utf-8-sig', 'gbk', 'gb18030']:
            try:
                with open(file, 'r', encoding=enc) as f:
                    for line in f:
                        line = line.strip()
                        if line:
                            accounts.add(line)
                break
            except:
                continue
        return accounts
    
    def filter_accounts(self):
        if not self.account_xlsx.get() or not self.account_txt.get():
            messagebox.showwarning("提示", "请选择两个文件")
            return
        
        self.clear_log2()
        bound = self.load_bound()
        all_acc = self.load_all()
        
        if not bound or not all_acc:
            return
        
        self.unbound_accounts = sorted(list(all_acc - bound))
        self.log2(f"已绑定：{len(bound)}")
        self.log2(f"总数：{len(all_acc)}")
        self.log2(f"未绑定：{len(self.unbound_accounts)}")
        self.log2("")
        for i, acc in enumerate(self.unbound_accounts, 1):
            self.log2(f"{i}. {acc}")
        
        self.stats_label.config(text=f"未绑定:{len(self.unbound_accounts)}")
        messagebox.showinfo("完成", f"未绑定：{len(self.unbound_accounts)} 个")
    
    def copy_accounts(self):
        if not self.unbound_accounts:
            messagebox.showwarning("提示", "请先筛选")
            return
        
        text = '\n'.join(self.unbound_accounts)
        self.root.clipboard_clear()
        self.root.clipboard_append(text)
        self.log2(f"已复制 {len(self.unbound_accounts)} 个ID")
        messagebox.showinfo("成功", f"已复制 {len(self.unbound_accounts)} 个ID")
    
    # ==================== 功能3：模板自动填写 ====================
    
    def setup_template_fill_tab(self):
        main_frame = ttk.Frame(self.tab3, padding="15", style='Card.TFrame')
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        self.saved_templates = []
        self.output_folder = tk.StringVar()
        
        # 所有占位符的变量
        self.company_var = tk.StringVar()  # {{company}}
        
        # 标题
        ttk.Label(main_frame, text="模板自动填写", style='Title.TLabel').pack(anchor='w', pady=(0, 5))
        
        # 提示
        if not HAS_DOCX:
            ttk.Label(main_frame, text="提示：需要安装python-docx库才能使用此功能\n运行: pip install python-docx", 
                     foreground='#E53935', background='#ffffff', font=('微软雅黑', 9)).pack(pady=5)
        
        # 第一部分：添加模板
        add_frame = ttk.LabelFrame(main_frame, text="添加模板", padding="10", style='Card.TLabelframe')
        add_frame.pack(fill=tk.X, pady=5)
        
        row1 = ttk.Frame(add_frame, style='Card.TFrame')
        row1.pack(fill=tk.X, pady=5)
        ttk.Label(row1, text="选择Word文件：", style='Card.TLabel').pack(side=tk.LEFT)
        self.new_template_path = tk.StringVar()
        e1 = tk.Entry(row1, textvariable=self.new_template_path, width=45, font=('微软雅黑', 9), relief='solid', bd=1)
        e1.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        e1.bind('<Double-Button-1>', lambda e: self.paste_path(self.new_template_path))
        self.make_draggable(e1, self.new_template_path, 'file', 'Word')
        ttk.Button(row1, text="选择文件", command=self.select_new_template, style='Secondary.TButton').pack(side=tk.LEFT, padx=2)
        ttk.Button(row1, text="添加并保存", command=self.add_template, style='Primary.TButton').pack(side=tk.LEFT, padx=2)
        
        # 第二部分：已保存的模板列表
        list_frame = ttk.LabelFrame(main_frame, text="已保存模板（勾选要生成的模板）", padding="10", style='Card.TLabelframe')
        list_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        self.template_list_frame = ttk.Frame(list_frame, style='Card.TFrame')
        self.template_list_frame.pack(fill=tk.BOTH, expand=True)
        
        btn_frame = ttk.Frame(list_frame, style='Card.TFrame')
        btn_frame.pack(fill=tk.X, pady=5)
        ttk.Button(btn_frame, text="全选", command=self.select_all_templates, style='Secondary.TButton').pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_frame, text="取消全选", command=self.unselect_all_templates, style='Secondary.TButton').pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_frame, text="删除选中", command=self.delete_selected_templates, style='Danger.TButton').pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_frame, text="扫描占位符", command=self.scan_placeholders_in_selected, style='Primary.TButton').pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_frame, text="批量更新路径", command=self.batch_update_paths, style='Secondary.TButton').pack(side=tk.LEFT, padx=2)
        
        # 第三部分：填写信息（动态显示）
        self.info_frame = ttk.LabelFrame(main_frame, text="填写信息（选择模板后点击'扫描占位符'显示）", padding="10", style='Card.TLabelframe')
        self.info_frame.pack(fill=tk.X, pady=5)
        
        # 创建所有可能的输入框（初始隐藏）
        self.input_rows = {}
        
        # {{company}} - 主体名称
        self.company_row = ttk.Frame(self.info_frame, style='Card.TFrame')
        ttk.Label(self.company_row, text="主体名称 {{company}}：", style='Card.TLabel').pack(side=tk.LEFT)
        tk.Entry(self.company_row, textvariable=self.company_var, width=30, font=('微软雅黑', 9), relief='solid', bd=1).pack(side=tk.LEFT, padx=5)
        
        # {{Date_of_Signing}} - 签订日期（自动计算）
        self.signing_row = ttk.Frame(self.info_frame, style='Card.TFrame')
        self.signing_date_var = tk.StringVar()
        ttk.Label(self.signing_row, text="签订日期 {{Date_of_Signing}}：", style='Card.TLabel').pack(side=tk.LEFT)
        ttk.Label(self.signing_row, textvariable=self.signing_date_var, width=20, style='Card.TLabel').pack(side=tk.LEFT, padx=5)
        ttk.Label(self.signing_row, text="(自动计算)", style='Hint.TLabel').pack(side=tk.LEFT)
        
        # {{out_of_date}} - 到期日（自动计算）
        self.expiry_row = ttk.Frame(self.info_frame, style='Card.TFrame')
        self.expiry_date_var = tk.StringVar()
        ttk.Label(self.expiry_row, text="到期日 {{out_of_date}}：", style='Card.TLabel').pack(side=tk.LEFT)
        ttk.Label(self.expiry_row, textvariable=self.expiry_date_var, width=20, style='Card.TLabel').pack(side=tk.LEFT, padx=5)
        ttk.Label(self.expiry_row, text="(签订日期+1年-1天)", style='Hint.TLabel').pack(side=tk.LEFT)
        
        # {{date}} - 普通日期
        self.date_row = ttk.Frame(self.info_frame, style='Card.TFrame')
        self.date_var = tk.StringVar()
        ttk.Label(self.date_row, text="日期 {{date}}：", style='Card.TLabel').pack(side=tk.LEFT)
        ttk.Label(self.date_row, textvariable=self.date_var, width=20, style='Card.TLabel').pack(side=tk.LEFT, padx=5)
        ttk.Label(self.date_row, text="(当天日期)", style='Hint.TLabel').pack(side=tk.LEFT)
        
        # 保存输入框引用
        self.input_rows = {
            'company': self.company_row,
            'Date_of_Signing': self.signing_row,
            'out_of_date': self.expiry_row,
            'date': self.date_row
        }
        
        # 初始隐藏所有输入框
        for row in self.input_rows.values():
            row.pack_forget()
        
        # 第四部分：输出设置
        output_frame = ttk.Frame(main_frame, style='Card.TFrame')
        output_frame.pack(fill=tk.X, pady=5)
        ttk.Label(output_frame, text="输出文件夹：", style='Card.TLabel', width=12).pack(side=tk.LEFT)
        e2 = tk.Entry(output_frame, textvariable=self.output_folder, width=45, font=('微软雅黑', 9), relief='solid', bd=1)
        e2.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        e2.bind('<Double-Button-1>', lambda e: self.paste_path(self.output_folder))
        self.make_draggable(e2, self.output_folder, 'folder', '输出文件夹')
        ttk.Button(output_frame, text="选择", command=self.select_output_folder, style='Secondary.TButton').pack(side=tk.LEFT, padx=2)
        
        # 第五部分：操作按钮
        row5 = ttk.Frame(main_frame, style='Card.TFrame')
        row5.pack(fill=tk.X, pady=10)
        ttk.Button(row5, text="📋 预览替换", command=self.preview_template, style='Secondary.TButton').pack(side=tk.LEFT, padx=5)
        ttk.Button(row5, text="✅ 生成合同", command=self.generate_contract, style='Primary.TButton').pack(side=tk.LEFT, padx=5)
        
        # 日志区域
        log_frame = ttk.LabelFrame(main_frame, text="操作日志", padding="10", style='Card.TLabelframe')
        log_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        self.log_text3 = ScrolledText(log_frame, height=8, font=('Consolas', 9), bg='#fafafa', relief='flat', bd=0)
        self.log_text3.pack(fill=tk.BOTH, expand=True)
        
        # 加载已保存的模板
        self.template_checkboxes = []  # 初始化复选框列表
        self.found_placeholders = set()  # 找到的占位符
        self.load_saved_templates()
        self.refresh_template_list()
    
    def get_signing_date(self):
        """计算签订日期：取小于等于当日的最近合规签约日（1日、15日、月末）"""
        today = datetime.date.today()
        day = today.day
        year = today.year
        month = today.month
        
        # 获取当月最后一天
        if month == 12:
            last_day = 31
        else:
            last_day = (datetime.date(year, month + 1, 1) - datetime.timedelta(days=1)).day
        
        # 合规签约日：1、15、月末
        valid_days = [1, 15, last_day]
        
        # 如果当天是合规日，直接用当天
        if day in valid_days:
            return today
        
        # 找小于等于当天的最近合规日
        candidates = [d for d in valid_days if d <= day]
        if candidates:
            target_day = max(candidates)
            return datetime.date(year, month, target_day)
        else:
            # 如果当天小于1，用上月月末（一般不会发生）
            return datetime.date(year, month, 1)
    
    def get_expiry_date(self, signing_date):
        """计算到期日：签订日期+1年-1天"""
        # 往后推1年
        try:
            one_year_later = datetime.date(signing_date.year + 1, signing_date.month, signing_date.day)
        except ValueError:
            # 处理2月29日等特殊情况
            one_year_later = datetime.date(signing_date.year + 1, signing_date.month, 28)
        
        # 减1天
        expiry_date = one_year_later - datetime.timedelta(days=1)
        return expiry_date
    
    def format_date_cn(self, date_obj):
        """格式化日期为中文格式"""
        return f"{date_obj.year}年{date_obj.month}月{date_obj.day}日"
    
    def scan_placeholders_in_selected(self):
        """扫描选中模板中的占位符，只显示存在的输入框"""
        if not HAS_DOCX:
            messagebox.showwarning("提示", "请先安装python-docx库")
            return
        
        # 获取选中的模板
        selected = []
        for var, template in self.template_checkboxes:
            if var.get():
                selected.append(template)
        
        if not selected:
            messagebox.showwarning("提示", "请先勾选模板")
            return
        
        # 扫描所有选中模板的占位符
        all_placeholders = set()
        
        for template in selected:
            path = template['path']
            if not os.path.exists(path):
                continue
            
            try:
                doc = Document(path)
                for para in doc.paragraphs:
                    text = para.text
                    for placeholder in ['{{company}}', '{{Date_of_Signing}}', '{{out_of_date}}', '{{date}}']:
                        if placeholder in text:
                            # 提取占位符名称
                            name = placeholder.replace('{{', '').replace('}}', '')
                            all_placeholders.add(name)
                
                # 检查表格
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                text = para.text
                                for placeholder in ['{{company}}', '{{Date_of_Signing}}', '{{out_of_date}}', '{{date}}']:
                                    if placeholder in text:
                                        name = placeholder.replace('{{', '').replace('}}', '')
                                        all_placeholders.add(name)
            except Exception as e:
                self.log3(f"扫描错误：{template['name']} - {e}")
        
        self.found_placeholders = all_placeholders
        
        # 隐藏所有输入框
        for row in self.input_rows.values():
            row.pack_forget()
        
        # 只显示存在的占位符输入框
        if 'company' in all_placeholders:
            self.company_row.pack(fill=tk.X, pady=5)
        
        if 'Date_of_Signing' in all_placeholders:
            signing_date = self.get_signing_date()
            self.signing_date_var.set(self.format_date_cn(signing_date))
            self.signing_row.pack(fill=tk.X, pady=5)
        
        if 'out_of_date' in all_placeholders:
            # 先计算签订日期
            signing_date = self.get_signing_date()
            expiry_date = self.get_expiry_date(signing_date)
            self.signing_date_var.set(self.format_date_cn(signing_date))
            self.expiry_date_var.set(self.format_date_cn(expiry_date))
            self.expiry_row.pack(fill=tk.X, pady=5)
            # 如果有签订日期，也显示它
            if 'Date_of_Signing' not in all_placeholders:
                self.signing_row.pack(fill=tk.X, pady=5)
        
        if 'date' in all_placeholders:
            self.date_var.set(self.format_date_cn(datetime.date.today()))
            self.date_row.pack(fill=tk.X, pady=5)
        
        self.log3(f"找到占位符：{', '.join(all_placeholders) if all_placeholders else '无'}")
        if all_placeholders:
            messagebox.showinfo("完成", f"找到占位符：{', '.join(all_placeholders)}\n已显示对应输入框")
        else:
            messagebox.showinfo("提示", "选中模板中未找到占位符")
    
    def log3(self, msg):
        self.log_text3.insert(tk.END, msg + "\n")
        self.log_text3.see(tk.END)
    
    def clear_log3(self):
        self.log_text3.delete(1.0, tk.END)
    
    def get_config_path(self):
        """获取配置文件路径"""
        # 配置文件保存在程序所在目录
        if hasattr(sys, 'frozen'):
            # 打包后的EXE
            base_dir = os.path.dirname(sys.executable)
        else:
            base_dir = os.path.dirname(os.path.abspath(__file__))
        return os.path.join(base_dir, TEMPLATE_CONFIG_FILE)
    
    def load_saved_templates(self):
        """加载已保存的模板配置"""
        try:
            config_path = self.get_config_path()
            if os.path.exists(config_path):
                with open(config_path, 'r', encoding='utf-8') as f:
                    self.saved_templates = json.load(f)
                self.log3(f"已加载 {len(self.saved_templates)} 个模板配置")
        except Exception as e:
            self.saved_templates = []
            self.log3(f"加载配置失败：{e}")
    
    def save_template_config(self):
        """保存模板配置"""
        try:
            config_path = self.get_config_path()
            with open(config_path, 'w', encoding='utf-8') as f:
                json.dump(self.saved_templates, f, ensure_ascii=False, indent=2)
            self.log3("配置已保存")
        except Exception as e:
            self.log3(f"保存配置失败：{e}")
    
    def refresh_template_list(self):
        """刷新模板列表显示"""
        # 清空现有列表
        for widget in self.template_list_frame.winfo_children():
            widget.destroy()
        
        self.template_checkboxes = []
        
        if not self.saved_templates:
            ttk.Label(self.template_list_frame, text="暂无保存的模板，请添加模板").pack(pady=20)
            return
        
        for i, template in enumerate(self.saved_templates):
            name = template.get('name', f'模板{i+1}')
            path = template.get('path', '')
            
            row = ttk.Frame(self.template_list_frame)
            row.pack(fill=tk.X, pady=2)
            
            var = tk.BooleanVar(value=False)
            
            # 检查文件是否存在
            file_exists = os.path.exists(path) if path else False
            
            # 使用tk.Checkbutton避免显示横杠问题
            if file_exists:
                cb = tk.Checkbutton(row, text=name, variable=var, anchor='w')
                cb.pack(side=tk.LEFT)
            else:
                cb = tk.Checkbutton(row, text=f"{name} (文件不存在)", variable=var, anchor='w')
                cb.pack(side=tk.LEFT)
                # 添加重新选择按钮
                ttk.Button(row, text="重新选择", width=10,
                          command=lambda idx=i: self.update_template_path(idx)).pack(side=tk.LEFT, padx=5)
            
            self.template_checkboxes.append((var, template))
    
    def update_template_path(self, index):
        """更新模板文件路径"""
        file = filedialog.askopenfilename(
            title="重新选择模板文件",
            filetypes=[("Word文档", "*.docx"), ("全部", "*.*")]
        )
        
        if file:
            if not file.endswith('.docx'):
                messagebox.showwarning("提示", "请选择.docx格式的文件")
                return
            
            self.saved_templates[index]['path'] = file
            self.save_template_config()
            self.refresh_template_list()
            self.log3(f"已更新模板路径：{self.saved_templates[index]['name']}")
            messagebox.showinfo("成功", "模板路径已更新")
    
    def select_new_template(self):
        """选择新的模板文件"""
        file = filedialog.askopenfilename(
            title="选择Word模板", 
            filetypes=[("Word文档", "*.docx"), ("全部", "*.*")]
        )
        if file:
            self.new_template_path.set(file)
            self.log3(f"已选择：{file}")
    
    def add_template(self):
        """添加并保存模板"""
        path = self.new_template_path.get()
        if not path:
            messagebox.showwarning("提示", "请先选择Word模板文件")
            return
        
        if not path.endswith('.docx'):
            messagebox.showwarning("提示", "请选择.docx格式的Word文件")
            return
        
        if not os.path.exists(path):
            messagebox.showwarning("提示", "文件不存在")
            return
        
        # 获取模板名称（文件名）
        name = os.path.basename(path).replace('.docx', '')
        
        # 添加到保存列表
        template_info = {
            'name': name,
            'path': path
        }
        
        # 检查是否已存在
        for t in self.saved_templates:
            if t['path'] == path:
                messagebox.showinfo("提示", "该模板已保存")
                return
        
        self.saved_templates.append(template_info)
        self.save_template_config()
        self.refresh_template_list()
        self.log3(f"已添加模板：{name}")
        self.new_template_path.set('')
        messagebox.showinfo("成功", f"模板 '{name}' 已保存")
    
    def select_all_templates(self):
        """全选模板"""
        for var, template in self.template_checkboxes:
            var.set(True)
    
    def unselect_all_templates(self):
        """取消全选"""
        for var, template in self.template_checkboxes:
            var.set(False)
    
    def delete_selected_templates(self):
        """删除选中的模板"""
        to_delete = []
        for var, template in self.template_checkboxes:
            if var.get():
                to_delete.append(template)
        
        if not to_delete:
            messagebox.showwarning("提示", "请先勾选要删除的模板")
            return
        
        if not messagebox.askyesno("确认", f"确定删除 {len(to_delete)} 个模板配置？\n（不会删除原始文件）"):
            return
        
        for t in to_delete:
            self.saved_templates.remove(t)
        
        self.save_template_config()
        self.refresh_template_list()
        self.log3(f"已删除 {len(to_delete)} 个模板配置")
    
    def batch_update_paths(self):
        """批量更新模板路径"""
        # 选择包含模板文件的文件夹
        folder = filedialog.askdirectory(title="选择包含模板文件的文件夹")
        if not folder:
            return
        
        # 获取文件夹中的所有docx文件
        docx_files = {}
        for file in os.listdir(folder):
            if file.endswith('.docx'):
                name = file.replace('.docx', '')
                docx_files[name] = os.path.join(folder, file)
        
        if not docx_files:
            messagebox.showwarning("提示", "文件夹中没有docx文件")
            return
        
        # 尝试匹配并更新路径
        updated = 0
        not_found = []
        
        for template in self.saved_templates:
            name = template['name']
            if name in docx_files:
                template['path'] = docx_files[name]
                updated += 1
            else:
                not_found.append(name)
        
        if updated > 0:
            self.save_template_config()
            self.refresh_template_list()
        
        # 显示结果
        msg = f"成功更新 {updated} 个模板路径"
        if not_found:
            msg += f"\n未找到匹配文件：{', '.join(not_found)}"
        
        self.log3(msg)
        messagebox.showinfo("完成", msg)
    
    def select_output_folder(self):
        """选择输出文件夹"""
        folder = filedialog.askdirectory(title="选择输出文件夹")
        if folder:
            self.output_folder.set(folder)
            self.log3(f"输出文件夹：{folder}")
    
    def replace_placeholders_in_doc(self, doc, replacements):
        """替换文档中的占位符，尽量保留原有格式
        replacements: dict 格式，如 {'company': '某公司', 'Date_of_Signing': '2026年7月1日'}
        """
        
        def replace_in_paragraph(para):
            """在段落中替换占位符"""
            full_text = para.text
            
            # 检查是否包含任何占位符
            placeholders = ['{{company}}', '{{Date_of_Signing}}', '{{out_of_date}}', '{{date}}']
            has_placeholder = any(p in full_text for p in placeholders)
            
            if not has_placeholder:
                return
            
            # 逐个run检查并替换
            for run in para.runs:
                run_text = run.text
                for placeholder, value in replacements.items():
                    full_placeholder = '{{' + placeholder + '}}'
                    if full_placeholder in run_text:
                        run.text = run_text.replace(full_placeholder, value)
            
            # 处理占位符跨越多个run的情况
            remaining_text = para.text
            for placeholder, value in replacements.items():
                full_placeholder = '{{' + placeholder + '}}'
                if full_placeholder in remaining_text:
                    # 占位符跨越了多个run
                    runs_text = ''.join([r.text for r in para.runs])
                    
                    if full_placeholder in runs_text:
                        for i, run in enumerate(para.runs):
                            if '{{' in run.text and '}}' not in run.text:
                                accumulated = run.text
                                end_idx = i
                                
                                for j in range(i + 1, len(para.runs)):
                                    accumulated += para.runs[j].text
                                    if '}}' in accumulated:
                                        end_idx = j
                                        break
                                
                                for ph, val in replacements.items():
                                    full_ph = '{{' + ph + '}}'
                                    if full_ph in accumulated:
                                        new_text = accumulated.replace(full_ph, val)
                                        run.text = new_text
                                        for j in range(i + 1, end_idx + 1):
                                            para.runs[j].text = ''
        
        # 替换段落中的占位符
        for para in doc.paragraphs:
            replace_in_paragraph(para)
        
        # 替换表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_in_paragraph(para)
    
    def preview_template(self):
        """预览替换"""
        if not HAS_DOCX:
            messagebox.showwarning("提示", "请先安装python-docx库")
            return
        
        # 获取选中的模板
        selected = []
        for var, template in self.template_checkboxes:
            if var.get():
                selected.append(template)
        
        if not selected:
            messagebox.showwarning("提示", "请先勾选要生成的模板")
            return
        
        # 构建替换字典
        replacements = self.build_replacements()
        
        self.clear_log3()
        self.log3("预览替换内容：")
        for ph, val in replacements.items():
            self.log3(f"  {{{{{ph}}}}} → {val}")
        self.log3("")
        self.log3(f"已选择 {len(selected)} 个模板：")
        
        for template in selected:
            path = template['path']
            name = template['name']
            
            if not os.path.exists(path):
                self.log3(f"  {name} - 文件不存在！")
                continue
            
            try:
                doc = Document(path)
                found = {}
                for para in doc.paragraphs:
                    for ph in ['company', 'Date_of_Signing', 'out_of_date', 'date']:
                        if '{{' + ph + '}}' in para.text:
                            found[ph] = True
                
                status = ', '.join([f"{ph}:有" for ph in found.keys()]) if found else "无占位符"
                self.log3(f"  {name} - {status}")
                
            except Exception as e:
                self.log3(f"  {name} - 错误：{e}")
    
    def build_replacements(self):
        """构建替换字典"""
        replacements = {}
        
        # company
        company = self.company_var.get()
        if company:
            replacements['company'] = company
        
        # 签订日期
        signing_date = self.get_signing_date()
        replacements['Date_of_Signing'] = self.format_date_cn(signing_date)
        
        # 到期日
        expiry_date = self.get_expiry_date(signing_date)
        replacements['out_of_date'] = self.format_date_cn(expiry_date)
        
        # 普通日期
        replacements['date'] = self.format_date_cn(datetime.date.today())
        
        return replacements
    
    def generate_contract(self):
        """生成合同"""
        if not HAS_DOCX:
            messagebox.showwarning("提示", "请先安装python-docx库")
            return
        
        # 获取选中的模板
        selected = []
        for var, template in self.template_checkboxes:
            if var.get():
                selected.append(template)
        
        if not selected:
            messagebox.showwarning("提示", "请先勾选要生成的模板")
            return
        
        output = self.output_folder.get()
        if not output:
            messagebox.showwarning("提示", "请选择输出文件夹")
            return
        
        # 检查company是否需要填写
        if 'company' in self.found_placeholders:
            company = self.company_var.get()
            if not company:
                messagebox.showwarning("提示", "请填写主体名称")
                return
        
        # 构建替换字典
        replacements = self.build_replacements()
        
        self.clear_log3()
        self.log3("开始生成合同...")
        for ph, val in replacements.items():
            self.log3(f"{ph}: {val}")
        self.log3("")
        
        ok, err = 0, 0
        
        for template in selected:
            path = template['path']
            name = template['name']
            
            try:
                if not os.path.exists(path):
                    self.log3(f"[失败] {name} - 文件不存在")
                    err += 1
                    continue
                
                doc = Document(path)
                self.replace_placeholders_in_doc(doc, replacements)
                
                # 生成新文件名
                company = self.company_var.get() if self.company_var.get() else "合同"
                new_name = f"{name}_{company}.docx"
                dst = os.path.join(output, new_name)
                doc.save(dst)
                self.log3(f"[成功] {name} → {new_name}")
                ok += 1
                
            except Exception as e:
                self.log3(f"[失败] {name}：{e}")
                err += 1
        
        self.log3("")
        self.log3(f"完成！成功 {ok}，失败 {err}")
        messagebox.showinfo("完成", f"成功生成 {ok} 个合同文件\n保存至：{output}")
    
    # ==================== 功能4：数值转换 ====================
    
    def setup_value_convert_tab(self):
        """设置数值转换标签页"""
        main_frame = ttk.Frame(self.tab4, padding="15", style='Card.TFrame')
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 标题
        ttk.Label(main_frame, text="数值转换", style='Title.TLabel').pack(anchor='w', pady=(0, 10))
        
        # 输入区域
        input_frame = ttk.LabelFrame(main_frame, text="输入ID", padding="10", style='Card.TLabelframe')
        input_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        self.input_text = tk.Text(input_frame, height=8, width=60, font=('Consolas', 10), relief='solid', bd=1, bg='#fafafa')
        self.input_text.pack(fill=tk.BOTH, expand=True)
        
        # 输出格式选择
        format_frame = ttk.Frame(main_frame, style='Card.TFrame')
        format_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(format_frame, text="输出格式：", style='Card.TLabel').pack(side=tk.LEFT)
        self.output_format = tk.StringVar(value="按行")
        ttk.Radiobutton(format_frame, text="按行分隔（所有ID一行）", variable=self.output_format, value="按行", style='Custom.TRadiobutton').pack(side=tk.LEFT, padx=10)
        ttk.Radiobutton(format_frame, text="按列分隔（每个ID一行）", variable=self.output_format, value="按列", style='Custom.TRadiobutton').pack(side=tk.LEFT, padx=10)
        
        # 分隔符选择
        separator_frame = ttk.Frame(main_frame, style='Card.TFrame')
        separator_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(separator_frame, text="分隔符：", style='Card.TLabel').pack(side=tk.LEFT)
        self.output_separator = tk.StringVar(value="逗号")
        ttk.Radiobutton(separator_frame, text="英文逗号（,）", variable=self.output_separator, value="逗号", style='Custom.TRadiobutton').pack(side=tk.LEFT, padx=10)
        ttk.Radiobutton(separator_frame, text="英文分号（;）", variable=self.output_separator, value="分号", style='Custom.TRadiobutton').pack(side=tk.LEFT, padx=10)
        
        # 操作按钮
        btn_frame = ttk.Frame(main_frame, style='Card.TFrame')
        btn_frame.pack(fill=tk.X, pady=10)
        
        ttk.Button(btn_frame, text="🔄 转换", command=self.convert_values, style='Primary.TButton').pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="📋 复制结果", command=self.copy_convert_result, style='Success.TButton').pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="🗑️ 清空", command=self.clear_convert, style='Danger.TButton').pack(side=tk.LEFT, padx=5)
        
        # 输出区域
        output_frame = ttk.LabelFrame(main_frame, text="转换结果", padding="10", style='Card.TLabelframe')
        output_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        self.output_text = tk.Text(output_frame, height=8, width=60, font=('Consolas', 10), relief='solid', bd=1, bg='#fafafa')
        self.output_text.pack(fill=tk.BOTH, expand=True)
    
    def convert_values(self):
        """执行数值转换"""
        input_content = self.input_text.get("1.0", tk.END).strip()
        if not input_content:
            messagebox.showwarning("提示", "请输入ID")
            return
        
        output_format = self.output_format.get()
        output_sep = self.output_separator.get()
        
        # 解析ID（自动处理多种输入格式）
        parts = input_content.replace('\n', ' ').replace(',', ' ').replace(';', ' ').split()
        ids = [p.strip() for p in parts if p.strip()]
        
        if not ids:
            messagebox.showwarning("提示", "未找到有效的ID")
            return
        
        # 分隔符
        sep = ',' if output_sep == "逗号" else ';'
        
        # 根据输出格式生成结果
        if output_format == "按行":
            # 所有ID一行，用分隔符连接
            result = sep.join(ids)
        else:
            # 每个ID一行，后面加分隔符
            result = '\n'.join([id + sep for id in ids])
        
        # 显示结果
        self.output_text.delete("1.0", tk.END)
        self.output_text.insert("1.0", result)
        
        messagebox.showinfo("完成", f"转换完成\n共 {len(ids)} 个ID")
    
    def copy_convert_result(self):
        """复制转换结果到剪贴板"""
        result = self.output_text.get("1.0", tk.END).strip()
        if not result:
            messagebox.showwarning("提示", "没有结果可复制")
            return
        
        self.root.clipboard_clear()
        self.root.clipboard_append(result)
        messagebox.showinfo("成功", "已复制到剪贴板")
    
    def clear_convert(self):
        """清空输入和输出"""
        self.input_text.delete("1.0", tk.END)
        self.output_text.delete("1.0", tk.END)

    # ==================== 功能5：表格匹配 ====================

    def setup_table_match_tab(self):
        """设置表格匹配标签页"""
        main_frame = ttk.Frame(self.tab5, padding="15", style='Card.TFrame')
        main_frame.pack(fill=tk.BOTH, expand=True)

        # 标题
        ttk.Label(main_frame, text="表格匹配工具", style='Title.TLabel').pack(anchor='w', pady=(0, 5))
        ttk.Label(main_frame, text="以关联键为匹配条件，将原始表格字段带出到待匹配表格", style='Hint.TLabel').pack(anchor='w', pady=(0, 10))

        # 原始表格区域
        row1 = ttk.Frame(main_frame, style='Card.TFrame')
        row1.pack(fill=tk.X, pady=5)
        ttk.Label(row1, text="原始表格：", style='Card.TLabel', width=12).pack(side=tk.LEFT)
        self.match_source_path = tk.StringVar()
        e1 = tk.Entry(row1, textvariable=self.match_source_path, width=55, font=('微软雅黑', 9), relief='solid', bd=1)
        e1.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        e1.bind('<Double-Button-1>', lambda e: self.paste_path(self.match_source_path))
        self.make_draggable(e1, self.match_source_path, 'file', '表格')
        ttk.Button(row1, text="选择", command=lambda: self.select_match_file('source'), style='Secondary.TButton').pack(side=tk.LEFT, padx=2)
        ttk.Button(row1, text="读取表头", command=lambda: self.read_match_headers('source'), style='Primary.TButton').pack(side=tk.LEFT, padx=2)
        self.source_header_label = ttk.Label(row1, text="未读取", style='Hint.TLabel')
        self.source_header_label.pack(side=tk.LEFT, padx=10)

        # 待匹配表格区域
        row2 = ttk.Frame(main_frame, style='Card.TFrame')
        row2.pack(fill=tk.X, pady=5)
        ttk.Label(row2, text="待匹配表格：", style='Card.TLabel', width=12).pack(side=tk.LEFT)
        self.match_target_path = tk.StringVar()
        e2 = tk.Entry(row2, textvariable=self.match_target_path, width=55, font=('微软雅黑', 9), relief='solid', bd=1)
        e2.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        e2.bind('<Double-Button-1>', lambda e: self.paste_path(self.match_target_path))
        self.make_draggable(e2, self.match_target_path, 'file', '表格')
        ttk.Button(row2, text="选择", command=lambda: self.select_match_file('target'), style='Secondary.TButton').pack(side=tk.LEFT, padx=2)
        ttk.Button(row2, text="读取表头", command=lambda: self.read_match_headers('target'), style='Primary.TButton').pack(side=tk.LEFT, padx=2)
        self.target_header_label = ttk.Label(row2, text="未读取", style='Hint.TLabel')
        self.target_header_label.pack(side=tk.LEFT, padx=10)

        # 关联键选择区域
        key_frame = ttk.LabelFrame(main_frame, text="关联键设置", padding="10", style='Card.TLabelframe')
        key_frame.pack(fill=tk.X, pady=5)

        key_row = ttk.Frame(key_frame, style='Card.TFrame')
        key_row.pack(fill=tk.X)
        ttk.Label(key_row, text="原始表关联键：", style='Card.TLabel').pack(side=tk.LEFT)
        self.source_key_var = tk.StringVar()
        self.source_key_combo = ttk.Combobox(key_row, textvariable=self.source_key_var, state='readonly', width=25)
        self.source_key_combo.pack(side=tk.LEFT, padx=5)

        ttk.Label(key_row, text="待匹配表关联键：", style='Card.TLabel').pack(side=tk.LEFT, padx=(20, 0))
        self.target_key_var = tk.StringVar()
        self.target_key_combo = ttk.Combobox(key_row, textvariable=self.target_key_var, state='readonly', width=25)
        self.target_key_combo.pack(side=tk.LEFT, padx=5)

        # 带出字段勾选区域
        field_frame = ttk.LabelFrame(main_frame, text="选择需要带出的字段", padding="10", style='Card.TLabelframe')
        field_frame.pack(fill=tk.BOTH, expand=True, pady=5)

        # 使用滚动区域放置复选框
        self.field_canvas = tk.Canvas(field_frame, bg='#ffffff', highlightthickness=0)
        field_scrollbar = ttk.Scrollbar(field_frame, orient='vertical', command=self.field_canvas.yview)
        self.field_inner = ttk.Frame(self.field_canvas, style='Card.TFrame')
        self.field_inner.bind('<Configure>', lambda e: self.field_canvas.configure(scrollregion=self.field_canvas.bbox('all')))
        self.field_canvas.create_window((0, 0), window=self.field_inner, anchor='nw')
        self.field_canvas.configure(yscrollcommand=field_scrollbar.set)
        self.field_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        field_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self.field_check_vars = {}  # 存储字段复选框变量
        self.source_headers = []     # 原始表头列表
        self.target_headers = []     # 待匹配表头列表
        self.match_hint_label = ttk.Label(self.field_inner, text="请先读取原始表格表头", style='Hint.TLabel')
        self.match_hint_label.pack(anchor='w')

        # 操作按钮区域
        btn_frame = ttk.Frame(main_frame, style='Card.TFrame')
        btn_frame.pack(fill=tk.X, pady=10)
        ttk.Button(btn_frame, text="✅ 执行匹配", command=self.execute_table_match, style='Success.TButton').pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="🗑️ 清空", command=self.clear_table_match, style='Danger.TButton').pack(side=tk.LEFT, padx=5)

        # 日志区域
        log_frame = ttk.LabelFrame(main_frame, text="操作日志", padding="10", style='Card.TLabelframe')
        log_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        self.log_text5 = ScrolledText(log_frame, height=8, font=('Consolas', 9), bg='#fafafa', relief='flat', bd=0)
        self.log_text5.pack(fill=tk.BOTH, expand=True)

    def log5(self, msg):
        self.log_text5.insert(tk.END, msg + "\n")
        self.log_text5.see(tk.END)

    def select_match_file(self, which):
        """选择表格文件"""
        file = filedialog.askopenfilename(
            title="选择表格文件",
            filetypes=[("Excel文件", "*.xlsx;*.xls"), ("CSV文件", "*.csv"), ("所有文件", "*.*")])
        if file:
            if which == 'source':
                self.match_source_path.set(file)
            else:
                self.match_target_path.set(file)

    def _read_table_headers(self, file_path):
        """读取表格表头，自动识别表头行（支持xlsx/xls/csv）"""
        if not file_path or not os.path.isfile(file_path):
            return None

        try:
            if file_path.lower().endswith('.csv'):
                import csv
                encodings = ['utf-8', 'gbk', 'gb18030', 'utf-8-sig']
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            # 只读前10行用于识别表头，加速读取
                            rows = []
                            for i, row in enumerate(csv.reader(f)):
                                rows.append(row)
                                if i >= 10:
                                    break
                        # 自动识别表头行
                        for row_idx, row in enumerate(rows):
                            non_empty = [str(c).strip() for c in row if str(c).strip()]
                            if len(non_empty) >= 2:
                                return [str(c).strip() for c in row], row_idx
                        return None
                    except:
                        continue
                return None
            else:
                # 使用 read_only 模式加速读取，只读前几行
                wb = load_workbook(file_path, read_only=True, data_only=True)
                ws = wb.active
                rows_iter = ws.iter_rows(min_row=1, max_row=5, values_only=True)
                for row_idx, row_values in enumerate(rows_iter, 1):
                    non_empty = [str(c).strip() for c in row_values if c is not None and str(c).strip()]
                    if len(non_empty) >= 2:
                        headers = [str(c).strip() if c is not None else '' for c in row_values]
                        wb.close()
                        return headers, row_idx - 1
                wb.close()
                return None
        except Exception as e:
            messagebox.showerror("错误", f"读取文件失败：{e}")
            return None

    def read_match_headers(self, which):
        """读取表头并更新下拉框"""
        if which == 'source':
            file_path = self.match_source_path.get()
            if not file_path:
                messagebox.showwarning("提示", "请先选择原始表格")
                return
            result = self._read_table_headers(file_path)
            if result is None:
                messagebox.showwarning("提示", "无法读取表头")
                return
            headers, header_row = result
            self.source_headers = headers
            self.source_key_combo['values'] = headers
            if headers:
                self.source_key_combo.current(0)
            self.source_header_label.config(text=f"表头在第{header_row + 1}行，共{len(headers)}列", style='Success.TLabel')
            self.log5(f"原始表格表头（第{header_row + 1}行）：{', '.join(headers)}")
            # 更新字段勾选区
            self._update_field_checkboxes()
        else:
            file_path = self.match_target_path.get()
            if not file_path:
                messagebox.showwarning("提示", "请先选择待匹配表格")
                return
            result = self._read_table_headers(file_path)
            if result is None:
                messagebox.showwarning("提示", "无法读取表头")
                return
            headers, header_row = result
            self.target_headers = headers
            self.target_key_combo['values'] = headers
            if headers:
                self.target_key_combo.current(0)
            self.target_header_label.config(text=f"表头在第{header_row + 1}行，共{len(headers)}列", style='Success.TLabel')
            self.log5(f"待匹配表格表头（第{header_row + 1}行）：{', '.join(headers)}")

    def _update_field_checkboxes(self):
        """更新字段勾选复选框"""
        # 清空现有复选框
        for widget in self.field_inner.winfo_children():
            widget.destroy()

        self.field_check_vars = {}

        if not self.source_headers:
            self.match_hint_label = ttk.Label(self.field_inner, text="请先读取原始表格表头", style='Hint.TLabel')
            self.match_hint_label.pack(anchor='w')
            return

        ttk.Label(self.field_inner, text="勾选需要从原始表格带出到待匹配表格的字段：", style='Card.TLabel').pack(anchor='w', pady=(0, 5))

        # 每行放4个复选框
        cols = 4
        row_frame = None
        for i, header in enumerate(self.source_headers):
            if i % cols == 0:
                row_frame = ttk.Frame(self.field_inner, style='Card.TFrame')
                row_frame.pack(fill=tk.X, pady=2)
            var = tk.BooleanVar(value=False)
            self.field_check_vars[header] = var
            cb = tk.Checkbutton(row_frame, text=header, variable=var,
                                bg='#ffffff', font=('微软雅黑', 9),
                                activebackground='#ffffff')
            cb.pack(side=tk.LEFT, padx=10)

    def execute_table_match(self):
        """执行表格匹配，生成新工作簿（不修改原文件）"""
        source_path = self.match_source_path.get()
        target_path = self.match_target_path.get()

        if not source_path or not os.path.isfile(source_path):
            messagebox.showwarning("提示", "请先选择原始表格")
            return
        if not target_path or not os.path.isfile(target_path):
            messagebox.showwarning("提示", "请先选择待匹配表格")
            return

        source_key = self.source_key_var.get()
        target_key = self.target_key_var.get()

        if not source_key:
            messagebox.showwarning("提示", "请选择原始表关联键")
            return
        if not target_key:
            messagebox.showwarning("提示", "请选择待匹配表关联键")
            return

        # 获取勾选的带出字段
        selected_fields = [h for h, v in self.field_check_vars.items() if v.get()]
        if not selected_fields:
            messagebox.showwarning("提示", "请至少勾选一个需要带出的字段")
            return

        # 选择保存路径
        save_path = filedialog.asksaveasfilename(
            title="保存匹配结果", defaultextension=".xlsx",
            filetypes=[("Excel", "*.xlsx")], initialfile="表格匹配结果.xlsx")
        if not save_path:
            return

        self.log5(f"\n开始匹配...")
        self.log5(f"原始表格：{source_path}")
        self.log5(f"待匹配表格：{target_path}")
        self.log5(f"关联键：{source_key} ↔ {target_key}")
        self.log5(f"带出字段：{', '.join(selected_fields)}")

        try:
            # 读取原始表格数据，构建关联键映射
            source_data = self._read_table_data(source_path, source_key, selected_fields)
            if not source_data:
                messagebox.showwarning("提示", "原始表格无有效数据")
                return

            self.log5(f"原始表格读取 {len(source_data)} 条记录")

            # 生成新工作簿：复制原始表、待匹配表，并生成匹配结果表
            matched_count = self._generate_match_result(
                source_path, target_path, save_path,
                target_key, selected_fields, source_data)

            self.log5(f"匹配完成，共匹配 {matched_count} 条")
            self.log5(f"结果已保存至：{save_path}")
            messagebox.showinfo("完成",
                f"匹配完成\n共匹配 {matched_count} 条\n\n"
                f"结果文件包含2个工作表：\n"
                f"1. 原始表格\n2. 匹配结果\n\n"
                f"文件：{save_path}")

        except Exception as e:
            messagebox.showerror("错误", str(e))
            self.log5(f"错误：{e}")

    def _read_table_data(self, file_path, key_field, selected_fields):
        """读取表格数据，返回 {关联键值: {字段: 值}} 字典"""
        result = {}

        if file_path.lower().endswith('.csv'):
            import csv
            encodings = ['utf-8', 'gbk', 'gb18030', 'utf-8-sig']
            rows = None
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        reader = csv.reader(f)
                        rows = list(reader)
                    break
                except:
                    continue
            if rows is None:
                return result

            # 找表头行
            header_row_idx = 0
            for row_idx, row in enumerate(rows):
                non_empty = [str(c).strip() for c in row if str(c).strip()]
                if len(non_empty) >= 2:
                    header_row_idx = row_idx
                    break

            headers = [str(c).strip() for c in rows[header_row_idx]]
            key_col = None
            field_cols = {}
            for ci, h in enumerate(headers):
                if h == key_field:
                    key_col = ci
                if h in selected_fields:
                    field_cols[h] = ci

            if key_col is None:
                return result

            for row in rows[header_row_idx + 1:]:
                if len(row) <= key_col:
                    continue
                key_val = str(row[key_col]).strip() if row[key_col] else ''
                if not key_val:
                    continue
                record = {}
                for field, col in field_cols.items():
                    record[field] = str(row[col]).strip() if col < len(row) and row[col] else ''
                result[key_val] = record
        else:
            wb = load_workbook(file_path, data_only=True)
            ws = wb.active

            # 找表头行
            header_row = 1
            for row_idx in range(1, min(ws.max_row + 1, 6)):
                row_values = [ws.cell(row=row_idx, column=c).value for c in range(1, ws.max_column + 1)]
                non_empty = [str(c).strip() for c in row_values if c is not None and str(c).strip()]
                if len(non_empty) >= 2:
                    header_row = row_idx
                    break

            headers = []
            for c in range(1, ws.max_column + 1):
                val = ws.cell(row=header_row, column=c).value
                headers.append(str(val).strip() if val else '')

            key_col = None
            field_cols = {}
            for ci, h in enumerate(headers):
                if h == key_field:
                    key_col = ci + 1
                if h in selected_fields:
                    field_cols[h] = ci + 1

            if key_col is None:
                wb.close()
                return result

            for r in range(header_row + 1, ws.max_row + 1):
                key_val = ws.cell(row=r, column=key_col).value
                key_val = str(key_val).strip() if key_val else ''
                if not key_val:
                    continue
                record = {}
                for field, col in field_cols.items():
                    val = ws.cell(row=r, column=col).value
                    record[field] = str(val).strip() if val else ''
                result[key_val] = record

            wb.close()

        return result

    def _read_table_all_rows(self, file_path):
        """读取表格所有行数据，返回 (表头行索引, 表头列表, 数据行列表)"""
        if file_path.lower().endswith('.csv'):
            import csv
            encodings = ['utf-8', 'gbk', 'gb18030', 'utf-8-sig']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        rows = list(csv.reader(f))
                    break
                except:
                    continue
            else:
                return None

            # 找表头行
            header_row_idx = 0
            for row_idx, row in enumerate(rows):
                non_empty = [str(c).strip() for c in row if str(c).strip()]
                if len(non_empty) >= 2:
                    header_row_idx = row_idx
                    break
            headers = [str(c).strip() for c in rows[header_row_idx]]
            data_rows = rows[header_row_idx + 1:]
            return header_row_idx, headers, data_rows
        else:
            wb = load_workbook(file_path, read_only=True, data_only=True)
            ws = wb.active
            all_rows = list(ws.iter_rows(values_only=True))
            wb.close()

            if not all_rows:
                return None

            # 找表头行
            header_row_idx = 0
            for row_idx, row in enumerate(all_rows):
                non_empty = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if len(non_empty) >= 2:
                    header_row_idx = row_idx
                    break

            headers = [str(c).strip() if c is not None else '' for c in all_rows[header_row_idx]]
            data_rows = []
            for row in all_rows[header_row_idx + 1:]:
                data_rows.append([str(c).strip() if c is not None else '' for c in row])
            return header_row_idx, headers, data_rows

    def _generate_match_result(self, source_path, target_path, save_path,
                                target_key, selected_fields, source_data):
        """生成新工作簿：原始表格 + 匹配结果，返回匹配数量"""
        # 读取原始表格所有行
        src_result = self._read_table_all_rows(source_path)
        # 读取待匹配表格所有行
        tgt_result = self._read_table_all_rows(target_path)

        if not src_result or not tgt_result:
            return 0

        _, src_headers, src_data_rows = src_result
        _, tgt_headers, tgt_data_rows = tgt_result

        # 找到待匹配表的关联键列索引
        key_col_idx = None
        for ci, h in enumerate(tgt_headers):
            if h == target_key:
                key_col_idx = ci
                break

        if key_col_idx is None:
            return 0

        # 构建匹配结果：待匹配表头 + 带出字段（新增的列）
        result_headers = list(tgt_headers)
        field_col_map = {}  # 字段 -> 结果表中的列索引
        for field in selected_fields:
            if field in result_headers:
                field_col_map[field] = result_headers.index(field)
            else:
                result_headers.append(field)
                field_col_map[field] = len(result_headers) - 1

        # 匹配并填充
        matched = 0
        result_rows = []
        for row in tgt_data_rows:
            # 复制行数据
            new_row = list(row)
            # 补齐列数
            while len(new_row) < len(result_headers):
                new_row.append('')

            # 获取关联键值
            key_val = str(row[key_col_idx]).strip() if key_col_idx < len(row) and row[key_col_idx] else ''
            if key_val and key_val in source_data:
                matched += 1
                for field in selected_fields:
                    new_row[field_col_map[field]] = source_data[key_val].get(field, '')

            result_rows.append(new_row)

        # 创建新工作簿
        wb = Workbook()

        # 工作表1：原始表格
        ws1 = wb.active
        ws1.title = "原始表格"
        ws1.append(src_headers)
        for row in src_data_rows:
            ws1.append(row)

        # 工作表2：匹配结果
        ws2 = wb.create_sheet(title="匹配结果")
        ws2.append(result_headers)
        for row in result_rows:
            ws2.append(row)

        # 设置列宽（简单自适应）
        for ws in [ws1, ws2]:
            for col_idx, header in enumerate(ws[1], 1):
                if header.value:
                    ws.column_dimensions[header.column_letter].width = min(max(len(str(header.value)) * 2, 12), 40)

        wb.save(save_path)
        wb.close()

        return matched

    def clear_table_match(self):
        """清空表格匹配区域"""
        self.match_source_path.set('')
        self.match_target_path.set('')
        self.source_headers = []
        self.target_headers = []
        self.source_key_combo['values'] = []
        self.target_key_combo['values'] = []
        self.source_key_var.set('')
        self.target_key_var.set('')
        self.source_header_label.config(text="未读取", style='Hint.TLabel')
        self.target_header_label.config(text="未读取", style='Hint.TLabel')
        self._update_field_checkboxes()
        self.log_text5.delete("1.0", tk.END)


import sys

def main():
    # 如果支持拖拽，使用 TkinterDnD.Tk
    if HAS_DND:
        root = TkinterDnD.Tk()
    else:
        root = tk.Tk()
    app = ToolboxApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()